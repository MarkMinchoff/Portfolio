# FILENAME: CV_GEN_V3.py
# AUTHOR: Mark Minchoff

# imports libraries
import PySimpleGUI as sg
import os
import sys
import docx
import docx2pdf
import time
from datetime import date
from docx.enum.text import WD_ALIGN_PARAGRAPH

# setting pysimplegui theme
sg.theme('System Default For Real')

# some cotainers for the primary code
mydoc = docx.Document()
today = date.today()
date = today.strftime("%m/%d/%y")

#########################################################################################

# layout for pysimplegui defined as a function
def MainGui():
    layout = [
    [sg.Text('Please Enter The Following Information', size=(70,1), justification = 'c')],
    [sg.Text('What Is Your Full Name? :', size =(30, 1)), sg.InputText()],
    [sg.Text('What Company Are You Applying To? :', size =(30, 1)), sg.InputText()],
    [sg.Text('What Position Are You Applying To? :', size =(30, 1)), sg.InputText()],
    [sg.Text('What Should We Call This Cover Letter? :', size =(30, 1)), sg.InputText()],
    # can't get this to work yet - I need to have user pick file save location
    # [sg.Text('Where Shall We Save Your Cover Letter? :', size =(30, 1)), sg.InputText(), sg.FolderBrowse(key="-IN-")],
    [sg.Button('Create'), sg.Exit(),sg.Text("                                                       "),
    sg.Checkbox('Create PDF Upon DOCX Completion?', default=False, key="-IN-")]
    ]

# title window name + container + layout connector

    window = sg.Window('Cover Letter Generator').Layout(layout)

#event loop start, keeps the window open
    while True:
        event, values = window.Read()
        # fail-safe to break or exit the event loop
        if event in (None, 'Exit'):
            break
        # if user hits create it runs the following code
        if event == 'Create':
            # setting containers for docx variables and adjustments
            save_loc = values["-IN-"]
            d = mydoc.add_paragraph(date)
            d.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            name = mydoc.add_paragraph(f'{values[0]}')
            name.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            # meat of the code to create docx document, add variables, and save
            p1 = mydoc.add_paragraph("Dear Hiring Manager,")
            p2 = mydoc.add_paragraph(f"I find myself to be a determined professional seeking the chance to succeed at {values[1]}.  I feel like I could be a good fit for the {values[2]} position.")
            p3 = mydoc.add_paragraph("If you are interested in my previous employment and would like me to elaborate my skill sets I will be more than happy to schedule a meeting with you. I am certain that I can be an asset in any position requiring hard work, enthusiasm and reliability.")
            p4 = mydoc.add_paragraph("I look forward to hearing from you. The enclosed resume lists all of my relevant experience and qualifications.")
            p5 = mydoc.add_paragraph("Due to their privacy I will only submit references once I am interviewed.")
            p6 = mydoc.add_paragraph("Thank you for your time and consideration.")
            p7 = mydoc.add_paragraph("Sincerely,")
            p9 = mydoc.add_paragraph(f"{values[0]}")
            mydoc.save(f'{values[3]}.docx')

            # if checkbox is selected then create a pdf of the cover letter
            if values["-IN-"] == True:
                sg.popup_auto_close('This Program Will Close Once Your PDF is Created')
                docx2pdf.convert(f'{values[3]}.docx')
                sys.exit()
            sys.exit()

# exit the window if loop is broken or 'exit' is selected
    window.close()

#########################################################################################

# you know what it is
if __name__ == '__main__':
    MainGui()
