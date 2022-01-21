# Cover Letter Generator v2.0
# Authored By Mark Minchoff

# import libraries
import os
import sys
import docx
import docx2pdf
import time
from datetime import date
from docx.enum.text import WD_ALIGN_PARAGRAPH

# setting containers for modules
mydoc = docx.Document()
today = date.today()
date = today.strftime("%m/%d/%y")

# kicks off the program
input("Press 'Enter' to Start")

# user selects their full name
u_name = input('What Is Your Full Name?:')
# user selects their potential company
company = input('What Company Are You Applying To?:')
# user selects their potential position
posi = input('What Position Are You Applying To?:')
 # user selects filename to save as without extension
cover = input('What Should We Call This Cover Letter?:')

# defined function, look into it
def C_Letter():
    d = mydoc.add_paragraph(date)
    d.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    name = mydoc.add_paragraph(f'{u_name}')
    name.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p1 = mydoc.add_paragraph("Dear Hiring Manager,")
    p2 = mydoc.add_paragraph(f"I find myself to be a determined professional seeking the chance to succeed at {company}.  I feel like I could be a good fit for the {posi} position.")
    p3 = mydoc.add_paragraph("If you are interested in my previous employment and would like me to elaborate my skill sets I will be more than happy to schedule a meeting with you. I am certain that I can be an asset in any position requiring hard work, enthusiasm and reliability.")
    p4 = mydoc.add_paragraph("I look forward to hearing from you. The enclosed resume lists all of my relevant experience and qualifications.")
    p5 = mydoc.add_paragraph("Due to their privacy I will only submit references once I am interviewed.")
    p6 = mydoc.add_paragraph("Thank you for your time and consideration.")
    p7 = mydoc.add_paragraph("Sincerely,")
    p9 = mydoc.add_paragraph(f"{u_name}")
    mydoc.save(f'{cover}.docx')
    time.sleep(1)
    print(f"{cover}.docx CREATION COMPLETE!\n")

# gives user the chance to bail out
quit = input('Do You Wish To Continue? Yes/No Only:')
if quit in ('no', 'NO', 'No', 'nO'):
    sys.exit()
# run the very important C_Letter function
else:
    C_Letter()

# Option to exit before turning into PDF
quit2 = input('Do You Wish To Create This Turn This Document Into A PDF? Yes/No Only:')
if quit2 in ('no', 'NO', 'No', 'nO'):
    sys.exit()

# wait 3 seconds and then turn cover letter into PDF
else:
    time.sleep(3)
    docx2pdf.convert(f'{cover}.docx')
    print('\n')
    print(f"{cover}.pdf CREATION COMPLETE!\n")
